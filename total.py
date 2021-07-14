from docx import Document
from copy import deepcopy
import sys,base64
from socket import inet_aton
import struct,os
def chk(item):
    for i in vulners:
        if item in i:
            return i
        else:
            return ''

all=open(sys.argv[1]).read()
per_host_result=all[all.index('II Results per Host')::]
hosts=all[all.index('Host Summary'):all.index('II Results per Host'):]
host_zone_donor=per_host_result.split('Port Summary for Host ')
host_zones=[]
host_zone_donor.remove(host_zone_donor[0])
for item in host_zone_donor:
    try:
        host_zones.append(item[:item.index('\n\nHost'):])
    except Exception,e:
        print e
        host_zones.append(item[::])
    

def do(host_zone):
    host_zone_Issues= host_zone.split('Issue\n-----')
    pre_host=host_zone_Issues[0]
    pre_pre_hostname=pre_host[pre_host.index('Security Issues for Host')+len('Security Issues for Host')::].strip()
    hostname=pre_pre_hostname.partition("\n")[0]
    host_zone_Issues.remove(host_zone_Issues[0])
    return {hostname:host_zone_Issues}

def pars_param(per_issue):
    per_issue_detailes=per_issue[:per_issue.index('\nSummary:\n'):]
    nvt=per_issue_detailes[per_issue_detailes.index('NVT:')+4:per_issue_detailes.index('\nOID:'):]
    threat_score=per_issue[per_issue.index('CVSS:')+5:per_issue.index(')\nPort:'):]
    threat_type=per_issue[per_issue.index('Threat:')+7:per_issue.index('(CVSS'):]
    port_pre=per_issue[per_issue.index('Port:')+8:per_issue.index('Port:')+28:]
    try:
        if 'tcp' in port_pre:
            port=port_pre[:port_pre.index('tcp'):]+'tcp'
        else:
            port=port_pre[:port_pre.index('udp'):]+'udp'
    except Exception,e:
        print e
    result={}
    if 'CVE:' in per_issue:
        pre_cve = per_issue[per_issue.index('CVE:') + 4::].strip()
        cve= pre_cve.partition("\n")[0]
        result.update({'cve':cve})
    else:
        result.update({'cve':'------'})
    if 'High' in threat_type.strip() and (float(threat_score.strip())>9):
        threat_type='Critical'
    result.update({'nvt': nvt.strip() , 'score':threat_score.strip(),'type':threat_type.strip(),'port': port.strip()})
    return result

def copy_table_after(ob_table, table):
    tbl, tb2 = ob_table._tbl, table._tbl
    new_tbl = deepcopy(tbl)
    tb2.addnext(new_tbl)

def copy_table_after3():
    p1, t1 = document.paragraphs[0]._p,document.tables[len(document.tables)-3]._tbl
    new_tbl = deepcopy(p1)
    t1.addnext(new_tbl)

def copy_table_after2(pragraphs, table):
    p, tb = pragraphs._p, table._tbl
    new_tbl = deepcopy(tb)
    p.addnext(new_tbl)

def row_count():
    target_table=document.tables[len(document.tables)-3]
    target_table.add_row()

def add_table():
    copy_table_after2(document.paragraphs[len(document.paragraphs)-3], document.tables[len(document.tables)-2])
    copy_table_after3()
    copy_table_after(document.tables[len(document.tables)-1], document.tables[len(document.tables)-3])
    #copy_table_after3()

def add_hostname(hostname):
    document.tables[len(document.tables)-4].cell(0, 0).text =hostname
    # table=document.tables[len(document.tables)-4]
    # cell = table.rows[0].cells[0]
    # paragraph = cell.text_frame.paragraphs[0]
    # paragraph.font.size = Pt(12)
    # paragraph.font.color.rgb = RGBColor(12, 34, 56)
    #document.tables[len(document.tables)-4].cell(0, 0).text =hostname
    document.tables[len(document.tables)-4].cell(3, 0).text =hostname

def import_vuln_count(crit,high,med,low):
    print '@@@@@@@@@@sssss@@@@@22'
    print crit
    print high
    print med
    print low
    print '@@@@@@@@@@@@@@@@@@@@@@22'
    tb_n=len(document.tables)-4
    tables=document.tables
    tables[tb_n].cell(3, 2).text = str(crit)
    tables[tb_n].cell(3, 3).text = str(high)
    tables[tb_n].cell(3, 4).text = str(med)
    tables[tb_n].cell(3, 5).text = str(low)

def insert_data(number,cve,nvt,port,threat_type):
    if 'general' in port:
        port='default'
    table=document.tables[len(document.tables)-3]
    row_s=table.rows
    table.cell(len(row_s)-1,4).text=number
    table.cell(len(row_s)-1,3).text=nvt
    table.cell(len(row_s)-1,1).text=port
    table.cell(len(row_s)-1,2).text=cve
    table.cell(len(row_s)-1,0).text=threat_type

def insert_parag(description):
    table=document.tables[len(document.tables)-2]
    tbl=  table._tbl
    tbl.addparagraph(description)
    #tbl.addnext(description)


document = Document('pure.docx')

list_host_zone={}
c=0
for i in host_zones:
    c=c+1
    print c
    targ= do(i)
    list_host_zone.update(targ)

vull=open('vuln.txt', 'w')
vull.close()
sorted_ip=sorted(list_host_zone.keys(), key=lambda ip: struct.unpack("!L", inet_aton(ip))[0])
critical={}
high={}
medium={}
low={}

bluekeep="Microsoft Windows Remote Desktop Services , Remote Code Execution Vulnerability (BlueKeep) - (Remote Active)"
ipblue=['10.60.2.6','10.60.3.45','10.60.3.46','10.60.3.49','10.60.4.8','10.60.4.18','10.60.4.30','10.60.4.51','10.60.4.55','10.60.5.9','10.60.7.8','10.60.7.20','10.60.7.241','10.60.9.10','10.60.9.201','10.60.9.215','10.60.9.216','10.60.4.40','10.60.4.44','10.60.6.11','10.60.6.12','10.60.6.37','10.60.6.132','10.60.7.10']

ipetrnal=['10.60.3.11','10.60.4.18','10.60.4.40','10.60.4.55','10.60.4.68','10.60.7.20','10.60.9.10','10.60.9.14','10.60.9.201','10.60.9.202','10.60.9.215']
etrnal="MS17-010: Security Update for Microsoft Windows SMB Server"

ipsmb=['10.60.9.14', '10.60.9.216', '10.60.9.201', '10.60.9.215', '10.60.7.20', '10.60.7.17', '10.60.7.10', '10.60.7.4', '10.60.7.5', '10.60.7.25', '10.60.7.7', '10.60.7.8', '10.60.7.21', '10.60.7.16', '10.60.6.134', '10.60.6.11', '10.60.6.7', '10.60.6.133', '10.60.6.6', '10.60.6.17', '10.60.6.19', '10.60.6.137', '10.60.6.10', '10.60.6.136', '10.60.6.132', '10.60.6.21', '10.60.6.37', '10.60.6.12', '10.60.5.7', '10.60.5.9', '10.60.5.4', '10.60.5.8', '10.60.4.5', '10.60.4.35', '10.60.4.53', '10.60.4.40', '10.60.4.42', '10.60.4.8', '10.60.4.18', '10.60.4.44', '10.60.4.59', '10.60.4.38', '10.60.4.33', '10.60.4.55', '10.60.4.51', '10.60.4.30', '10.60.4.7', '10.60.3.8', '10.60.3.20', '10.60.3.63', '10.60.3.15', '10.60.3.19', '10.60.3.62', '10.60.3.7', '10.60.3.17', '10.60.3.46', '10.60.3.61', '10.60.3.42', '10.60.3.16', '10.60.3.13', '10.60.3.49', '10.60.3.45', '10.60.3.60', '10.60.2.9', '10.60.2.6', '10.60.2.5', '10.60.2.22', '10.60.1.9', '10.60.1.137', '10.60.1.28', '10.60.1.135', '10.60.1.26', '10.60.1.23']
smbdef='Enabled (C$,IPC$,ADMIN$,etc) Policy and File Sharing miss configuration '

try:

    for item in sorted_ip:
        hname= item
        add_table()
        add_hostname(hname)
        count=1
        #---------------------should optimize --------------------------------------
        # ^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^smb ^^^^^^^^^^^^^^^^^^^66666
        if hname in ipsmb:
            row_count()
            insert_data(str(count), '', smbdef, '445', 'Critical')
            try:
                vuln_lines = open('vuln.txt', 'r').readlines()
                if not (smbdef + '\n' in vuln_lines):
                    vull = open('vuln.txt', 'a')
                    vull.write(smbdef)
                    vull.write('\n')
                    vull.close()
            except Exception, e:
                print e
            if smbdef in critical.keys():
                critical[smbdef] += 1
            else:
                critical.update({smbdef: 1})
            count += 1
        #^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^blue keep ^^^^^^^^^^^^^^^^^^^66666
        if hname in ipblue:
            row_count()
            insert_data(str(count), 'CVE-2019-0708', bluekeep, '3389', 'Critical')
            try:
                vuln_lines = open('vuln.txt', 'r').readlines()
                if not (bluekeep + '\n' in vuln_lines):
                    vull = open('vuln.txt', 'a')
                    vull.write(bluekeep)
                    vull.write('\n')
                    vull.close()
            except Exception, e:
                print e
            if bluekeep in critical.keys():
                critical[bluekeep] += 1
            else:
                critical.update({bluekeep: 1})
            count += 1

        #^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^eternal blue
        if hname in ipetrnal:
            row_count()
            insert_data(str(count), 'CVE-2017-0144', etrnal, '445', 'Critical')
            try:
                vuln_lines = open('vuln.txt', 'r').readlines()
                if not (bluekeep + '\n' in vuln_lines):
                    vull = open('vuln.txt', 'a')
                    vull.write(etrnal)
                    vull.write('\n')
                    vull.close()
            except Exception, e:
                print e
            if etrnal in critical.keys():
                critical[etrnal] += 1
            else:
                critical.update({etrnal: 1})
            count += 1
        #--------------------------------------------- end optimize --------------------------------
        cric_c = 0
        high_c = 0
        med_c = 0
        low_c = 0
        print '%%%%%%%%%%%%'+str(len(document.tables))
        #666666666666666666666666666666666
        for n in list_host_zone[item]:
                result=pars_param(n)
                if 'Critical' in result['type']:
                    cric_c+=1
                if 'High' in result['type']:
                    high_c +=1
                if 'Medium' in result['type']:
                    med_c +=1
                if 'Low' in result['type']:
                    low_c +=1
        import_vuln_count(cric_c,high_c,med_c,low_c)
        #78888888888888888888888888888888888
        for n in list_host_zone[item]:
            #if not('DCE/RPC and MSRPC Services Enumeration Reporting' in n or 'Missing `httpOnly` Cookie Attribute' in n):
            if True:
                row_count()
                result=pars_param(n)
                if 'Critical' in result['type']:
                    cric_c+=1
                    if result['nvt'] in critical.keys():
                        critical[result['nvt']]+=1
                    else:
                        critical.update({result['nvt']:1})
                if 'High' in result['type']:
                    high_c +=1
                    if result['nvt'] in high.keys():
                        high[result['nvt']] += 1
                    else:
                        high.update({result['nvt']: 1})

                if 'Medium' in result['type']:
                    med_c +=1
                    if result['nvt'] in medium.keys():
                        medium[result['nvt']] += 1
                    else:
                        medium.update({result['nvt']: 1})
                if 'Low' in result['type']:
                    low_c +=1
                    if result['nvt'] in low.keys():
                        low[result['nvt']] += 1
                    else:
                        low.update({result['nvt']: 1})

                insert_data(str(count), result['cve'], result['nvt'], result['port'], result['type'])
                try:
                    vuln_lines = open('vuln.txt', 'r').readlines()
                    if not (result['nvt'] + '\n' in vuln_lines):
                        vull = open('vuln.txt', 'a')
                        vull.write(result['nvt'])
                        vull.write('\n')
                        vull.close()
                except Exception, e:
                    print e
                count+=1

    file= sys.argv[1]
    document.save('Report_%s.docx'%file.replace('.txt',''))
    document_table = Document('table.docx')
except:
    document.save('%s.docx'%sys.argv[1])
