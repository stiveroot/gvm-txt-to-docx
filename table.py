from docx import Document
from copy import deepcopy
import sys,base64
from socket import inet_aton
import struct,os

def copy_table_after_table(ob_table, table):
    tbl, tb2 = ob_table._tbl, table._tbl
    new_tbl = deepcopy(tbl)
    tb2.addnext(new_tbl)

def copy_paragraph_after_table(document):
    p1, t1 = document.paragraphs[0]._p,document.tables[len(document.tables)-3]._tbl
    new_tbl = deepcopy(p1)
    t1.addnext(new_tbl)

def copy_table_after_paragraph(pragraphs, table):
    p, tb = pragraphs._p, table._tbl
    new_tbl = deepcopy(tb)
    p.addnext(new_tbl)

def add_row(document):
    target_table=document.tables[len(document.tables)-3]
    copy_table_after_table(document.tables[len(document.tables)-2], target_table)

def add_table(document):
    copy_table_after_paragraph(document.paragraphs[len(document.paragraphs)-3], document.tables[len(document.tables)-1])
    copy_paragraph_after_table(document)

def insert_data(nvt,number,threat_type,index,document):
    table=document.tables[len(document.tables)-3]
    row_s=table.rows
    table.cell(len(row_s)-1,3).text=str(number)
    table.cell(len(row_s)-1,1).text=nvt
    table.cell(len(row_s)-1,2).text=threat_type
    table.cell(len(row_s)-1,0).text=str(index)

def insert_parag(description):
    table=document.tables[len(document.tables)-3]
    tbl=  table._tbl
    tbl.addparagraph(description)




# critical= {'Microsoft SQL Server End Of Life Detection': 3, 'Microsoft Windows SMB Server Multiple Vulnerabilities-Remote (4013389)': 10, 'OS End Of Life Detection': 1}
# high= {'HiSilicon ASIC Firmware Multiple Vulnerabilities': 1, 'SMB Brute Force Logins With Default Credentials': 6, 'Generic HTTP Directory Traversal': 1, 'Deprecated SSH-1 Protocol Detection': 1}
# medium= {'SSL/TLS: Deprecated SSLv2 and SSLv3 Protocol Detection': 2, 'SSL/TLS: Report Vulnerable Cipher Suites for HTTPS': 1, 'SSL/TLS: Certificate Signed Using A Weak Signature Algorithm': 8, 'SSL/TLS: SSLv3 Protocol CBC Cipher Suites Information Disclosure Vulnerability (POODLE)': 2, 'Cleartext Transmission of Sensitive Information via HTTP': 1, 'SSL/TLS: Report Weak Cipher Suites': 10, 'SSL/TLS: Diffie-Hellman Key Exchange Insufficient DH Group Strength Vulnerability': 1, 'Telnet Unencrypted Cleartext Login': 4, 'JRun directory traversal': 1, 'SSH Weak Encryption Algorithms Supported': 1, 'SSL/TLS: OpenSSL CCS Man in the Middle Security Bypass Vulnerability': 1, "SSL/TLS: OpenSSL TLS 'heartbeat' Extension Information Disclosure Vulnerability": 1, 'FTP Unencrypted Cleartext Login': 3}
# low= {'SSH Weak MAC Algorithms Supported': 1, 'TCP timestamps': 71}


def iport_to_table(critical,high,medium,low,document,name):
    add_table(document)
    count=1
    for item in critical.keys():
        add_row(document)
        insert_data(item,critical[item],'crititcal',count,document)
        count+=1
    for item in high.keys():
        add_row(document)
        insert_data(item,high[item],'high',count,document)
        count+=1

    for item in medium.keys():
        add_row(document)
        insert_data(item,medium[item],'medium',count,document)
        count+=1

    for item in low.keys():
        add_row(document)
        insert_data(item,low[item],'low',count,document)
        count+=1
    document.save('total_vulner_%s_.docx'%name)

